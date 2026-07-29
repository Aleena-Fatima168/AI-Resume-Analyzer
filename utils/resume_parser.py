from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import BinaryIO, Union

import pdfplumber
import spacy
from docx import Document
from PyPDF2 import PdfReader

from config import ALLOWED_FILE_TYPES
from utils.text_processing import (
    clean_text,
    extract_section_lines,
    normalize_whitespace,
    split_into_lines,
)

_NLP: spacy.language.Language | None = None


def _get_nlp() -> spacy.language.Language | None:
    global _NLP
    if _NLP is None:
        try:
            _NLP = spacy.load("en_core_web_sm")
        except OSError:
            _NLP = None
    return _NLP


SUPPORTED_EXTENSIONS: tuple[str, ...] = tuple(sorted({*ALLOWED_FILE_TYPES, ".txt"}))

SourceType = Union[str, Path, BinaryIO]


@dataclass
class ResumeData:
    name:           str | None  = None
    email:          str | None  = None
    phone:          str | None  = None
    github:         str | None  = None
    linkedin:       str | None  = None
    education:      list[str]   = field(default_factory=list)
    experience:     list[str]   = field(default_factory=list)
    projects:       list[str]   = field(default_factory=list)
    certifications: list[str]   = field(default_factory=list)
    languages:      list[str]   = field(default_factory=list)
    skills:         list[str]   = field(default_factory=list)
    raw_text:       str         = ""

    def to_dict(self) -> dict:
        return {
            "name":           self.name,
            "email":          self.email,
            "phone":          self.phone,
            "github":         self.github,
            "linkedin":       self.linkedin,
            "education":      self.education,
            "experience":     self.experience,
            "projects":       self.projects,
            "certifications": self.certifications,
            "languages":      self.languages,
            "skills":         self.skills,
            "raw_text":       self.raw_text,
        }


class TextExtractor:
    def extract(self, source: SourceType, extension: str) -> str:
        raw = self._read_bytes(source)
        if extension == ".pdf":
            text = self._from_pdf(raw)
        elif extension == ".docx":
            text = self._from_docx(raw)
        elif extension == ".txt":
            text = self._from_txt(raw)
        else:
            raise ValueError(f"Unsupported extension: {extension}")
        return normalize_whitespace(clean_text(text))

    @staticmethod
    def _read_bytes(source: SourceType) -> bytes:
        if isinstance(source, (str, Path)):
            return Path(source).read_bytes()
        pos = source.tell() if hasattr(source, "tell") else None
        data = source.read()
        if pos is not None and hasattr(source, "seek"):
            source.seek(pos)
        return data

    def _from_pdf(self, data: bytes) -> str:
        text = self._pdf_pdfplumber(data)
        if not text.strip():
            text = self._pdf_pypdf2(data)
        if not text.strip():
            raise RuntimeError(
                "Could not extract text from PDF. "
                "The file may be scanned or image-based."
            )
        return text

    @staticmethod
    def _pdf_pdfplumber(data: bytes) -> str:
        pages: list[str] = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text(x_tolerance=2, y_tolerance=2)
                if page_text:
                    pages.append(page_text)
        return "\n".join(pages)

    @staticmethod
    def _pdf_pypdf2(data: bytes) -> str:
        reader = PdfReader(io.BytesIO(data))
        pages: list[str] = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                pages.append(page_text)
        return "\n".join(pages)

    @staticmethod
    def _from_docx(data: bytes) -> str:
        doc = Document(io.BytesIO(data))
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts)

    @staticmethod
    def _from_txt(data: bytes) -> str:
        for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
            try:
                return data.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                continue
        return data.decode("utf-8", errors="replace")


class ContactExtractor:
    _EMAIL    = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
    _PHONE    = re.compile(
        r"(?<!\d)"
        r"(\+?1[\s.\-]?)?"
        r"\(?\d{3}\)?[\s.\-]?\d{3}[\s.\-]?\d{4}"
        r"(?!\d)"
    )
    _GITHUB   = re.compile(
        r"(?:https?://)?(?:www\.)?github\.com/([A-Za-z0-9_\-]+)",
        re.IGNORECASE,
    )
    _LINKEDIN = re.compile(
        r"(?:https?://)?(?:www\.)?linkedin\.com/in/([A-Za-z0-9_\-]+)",
        re.IGNORECASE,
    )

    def extract_email(self, text: str) -> str | None:
        m = self._EMAIL.search(text)
        return m.group(0).lower() if m else None

    def extract_phone(self, text: str) -> str | None:
        m = self._PHONE.search(text)
        if not m:
            return None
        digits = re.sub(r"\D", "", m.group(0))
        if len(digits) == 10:
            return f"({digits[:3]}) {digits[3:6]}-{digits[6:]}"
        if len(digits) == 11 and digits[0] == "1":
            return f"+1 ({digits[1:4]}) {digits[4:7]}-{digits[7:]}"
        return m.group(0).strip()

    def extract_github(self, text: str) -> str | None:
        m = self._GITHUB.search(text)
        return f"https://github.com/{m.group(1)}" if m else None

    def extract_linkedin(self, text: str) -> str | None:
        m = self._LINKEDIN.search(text)
        return f"https://linkedin.com/in/{m.group(1)}" if m else None


class NameExtractor:
    _SKIP = re.compile(
        r"(resume|curriculum vitae|cv|@|http|linkedin|github|phone|email|address)",
        re.IGNORECASE,
    )
    _NAME_PATTERN = re.compile(r"^[A-Z][a-zA-Z'\-]+(?:\s+[A-Z][a-zA-Z'\-]+){1,3}$")

    def extract(self, text: str) -> str | None:
        nlp = _get_nlp()
        if nlp is not None:
            doc = nlp(text[:500])
            for ent in doc.ents:
                if ent.label_ == "PERSON" and len(ent.text.split()) >= 2:
                    return ent.text.strip()

        for line in split_into_lines(text)[:15]:
            if self._SKIP.search(line):
                continue
            if self._NAME_PATTERN.match(line.strip()) and len(line.split()) <= 5:
                return line.strip()

        return None


class SectionExtractor:
    def extract(self, lines: list[str], section: str) -> list[str]:
        return extract_section_lines(lines, section)


class ResumeParser:
    def __init__(self, source: SourceType, filename: str | None = None) -> None:
        self._source    = source
        self._filename  = self._resolve_filename(source, filename)
        self._extension = self._resolve_extension(self._filename)

        if self._extension not in SUPPORTED_EXTENSIONS:
            allowed = ", ".join(e.lstrip(".").upper() for e in SUPPORTED_EXTENSIONS)
            raise ValueError(
                f"Unsupported file type '{self._extension}'. Allowed: {allowed}."
            )
        if isinstance(source, (str, Path)) and not Path(source).is_file():
            raise FileNotFoundError(f"Resume file not found: {source}")

        self._text_extractor    = TextExtractor()
        self._contact_extractor = ContactExtractor()
        self._name_extractor    = NameExtractor()
        self._section_extractor = SectionExtractor()

    @property
    def filename(self) -> str:
        return self._filename

    @property
    def extension(self) -> str:
        return self._extension

    @staticmethod
    def supported_extensions() -> tuple[str, ...]:
        return SUPPORTED_EXTENSIONS

    def parse(self) -> ResumeData:
        raw_text = self._text_extractor.extract(self._source, self._extension)
        lines    = split_into_lines(raw_text)

        data = ResumeData(raw_text=raw_text)
        data.name     = self._name_extractor.extract(raw_text)
        data.email    = self._contact_extractor.extract_email(raw_text)
        data.phone    = self._contact_extractor.extract_phone(raw_text)
        data.github   = self._contact_extractor.extract_github(raw_text)
        data.linkedin = self._contact_extractor.extract_linkedin(raw_text)

        data.education      = self._section_extractor.extract(lines, "education")
        data.experience     = self._section_extractor.extract(lines, "experience")
        data.projects       = self._section_extractor.extract(lines, "projects")
        data.certifications = self._section_extractor.extract(lines, "certifications")
        data.languages      = self._section_extractor.extract(lines, "languages")
        data.skills         = self._section_extractor.extract(lines, "skills")

        return data

    @staticmethod
    def _resolve_filename(source: SourceType, filename: str | None) -> str:
        if isinstance(source, (str, Path)):
            return Path(source).name
        if not filename:
            raise ValueError("filename is required when source is a binary stream.")
        return filename

    @staticmethod
    def _resolve_extension(name: str) -> str:
        suffix = Path(name).suffix.lower()
        if not suffix:
            raise ValueError(f"Cannot determine file extension from: '{name}'")
        return suffix
